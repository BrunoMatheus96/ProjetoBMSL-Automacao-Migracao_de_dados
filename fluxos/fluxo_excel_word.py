import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
import pandas as pd
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class excel_word:
    def __init__(self):
        self.file_path = ""

    def selec_planilha(self):
        # Cria uma janela “root” escondida
        root = tk.Tk()
        root.withdraw()  # Esconde a janela principal

        # Abre a janela para selecionar um arquivo
        self.file_path = filedialog.askopenfilename(
            title="Selecione um arquivo", filetypes=[("Todos os arquivos", "*.*")]
        )

        if self.file_path == "":
            print("Nenhum arquivo selecionado.")
        elif self.file_path.lower().endswith((".xlsx", ".xls")):
            print("Arquivo Excel selecionado:", self.file_path)
        else:
            print("O arquivo precisa ser do tipo .xlsx ou .xls.")

    def criar_doc(self, dados, base_path, planilha):
        try:

            def preparar_dados():
                df = pd.read_excel(self.file_path, sheet_name=planilha)
                df.columns = df.columns.str.strip()

                colunas = [
                    "Import Test Case Name",
                    "Test Case Target",
                    "Card de Desenvolvimento",
                    "Sistema",
                    "Responsável",
                    "Link do Card de Desenvolvimento",
                ]

                df_tabela = df[colunas].copy()
                df_tabela = df_tabela.dropna(how="all")
                df_tabela = df_tabela.apply(
                    lambda col: col.apply(
                        lambda x: x.strip() if isinstance(x, str) else x
                    )
                )

                return df, df_tabela


            def preparar_dados_steps():
                df = pd.read_excel(self.file_path, sheet_name="Planejamento")
                df.columns = df.columns.str.strip()

                colunas = [
                    "ID",
                    "Sistema",
                    "Step Name",
                    "Action",
                    "Expected Result",
                ]

                df_steps = df[colunas].copy()
                df_steps = df_steps.dropna(how="all")

                df_steps = df_steps.apply(
                    lambda col: col.apply(
                        lambda x: x.strip() if isinstance(x, str) else x
                    )
                )

                df_steps["ID"] = df_steps["ID"].ffill()
                df_steps["ID"] = pd.to_numeric(
                    df_steps["ID"], errors="coerce"
                ).astype("Int64")

                return df_steps


            def criar_um_documento(df, row):

                from docx import Document

                id_valor = df.loc[row.name, "ID"]
                id_str = "" if pd.isna(id_valor) else str(int(id_valor))

                doc = Document()

                header = doc.sections[0].header
                header_para = header.paragraphs[0]

                header_para.add_run().add_picture(
                    "imagens/minsait.png", width=Inches(2.20)
                )

                header_para.add_run(" " * 30)

                header_para.add_run().add_picture(
                    "imagens/petrobras.png", width=Inches(2.1)
                )

                title_para = header.add_paragraph()
                title_run = title_para.add_run("\nDocumento de Evidência")
                title_run.bold = True
                title_run.font.size = Pt(20)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                line_para = header.add_paragraph()
                run = line_para.add_run("_" * 100)
                run.bold = True
                line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                table = doc.add_table(rows=0, cols=2)
                table.style = "Table Grid"

                for c in row.index:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(c)
                    row_cells[1].text = str(row[c])

                for chave, valor in dados.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(chave)
                    row_cells[1].text = str(valor)

                return doc, id_str


            def adicionar_steps_doc(doc, df_steps, id_str):

                steps = df_steps[df_steps["ID"] == int(id_str)]

                if steps.empty:
                    return doc

                doc.add_page_break()

                doc.add_heading(level=1)

                for _, step in steps.iterrows():

                    p = doc.add_paragraph()

                    p.add_run(f"{step['Step Name']}: ").bold = True

                    p.add_run("\nAção: ").bold = True
                    p.add_run(f"{step['Action']}\n")

                    p.add_run("\nResultado Esperado: ").bold = True
                    p.add_run(f"{step['Expected Result']}\n\n")

                return doc


            df, df_tabela = preparar_dados()
            df_steps = preparar_dados_steps()

            os.makedirs(base_path, exist_ok=True)

            for _, row in df_tabela.iterrows():

                id_valor = df.loc[row.name, "ID"]
                id_str = "" if pd.isna(id_valor) else str(int(id_valor))

                nome_arquivo = f"Documento de teste - {row['Sistema']} - TC{id_str}.docx"
                full_path = os.path.join(base_path, nome_arquivo)

                # 🔎 Se arquivo já existe não faz nada
                if os.path.exists(full_path):
                    print(f"Arquivo '{nome_arquivo}' já existe")
                    continue

                # cria documento
                doc, id_str = criar_um_documento(df, row)

                # adiciona steps
                doc = adicionar_steps_doc(doc, df_steps, id_str)

                # salva documento
                doc.save(full_path)

                print(f"Documento gerado: {nome_arquivo}")

        except Exception as e:
            print("Erro ao criar documentos:", e)