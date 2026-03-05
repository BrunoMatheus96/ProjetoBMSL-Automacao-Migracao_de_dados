import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
import pandas as pd
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class excel_word:
    def __init__(self):
        self.file_path = ""

    def selec_planilha(self):
        # Cria uma janela “root” escondida
        root = tk.Tk()
        root.withdraw()  # Esconde a janela principal

        # Abre a janela para selecionar um arquivo
        self.file_path = filedialog.askopenfilename(
            title="Selecione um arquivo", filetypes=[("Todos os arquivos", "*.*")]
        )

        if self.file_path == "":
            print("Nenhum arquivo selecionado.")
        elif self.file_path.lower().endswith((".xlsx", ".xls")):
            print("Arquivo Excel selecionado:", self.file_path)
        else:
            print("O arquivo precisa ser do tipo .xlsx ou .xls.")

    def criar_doc(self, dados, base_path):
        try:

            def preparar_dados():
                df = pd.read_excel(self.file_path, sheet_name="Planejamento")
                df.columns = df.columns.str.strip()

                colunas = [
                    "Import Test Case Name",
                    "Test Case Target",
                    "Card de Desenvolvimento",
                    "Sistema",
                    "Responsável",
                    "Link do Card de Desenvolvimento",
                ]

                df_tabela = df[colunas].copy()
                df_tabela = df_tabela.dropna(how="all")
                df_tabela = df_tabela.apply(
                    lambda col: col.apply(
                        lambda x: x.strip() if isinstance(x, str) else x
                    )
                )

                return df, df_tabela

            # chamando a função interna
            df, df_tabela = preparar_dados()

            def criar_um_documento(df, row):
                # função que cria apenas um doc a partir de uma linha
                from docx import Document

                id_valor = df.loc[row.name, "ID"]
                id_str = "" if pd.isna(id_valor) else str(int(id_valor))

                # Cria o documento
                doc = Document()

                # Acessa o cabeçalho da primeira seção
                header = doc.sections[0].header
                # Adiciona um parágrafo ao cabeçalho
                header_para = header.paragraphs[0]

                # Adiciona imagens ao cabeçalho
                run = header_para.add_run()
                header_para.add_run().add_picture(
                    "imagens/minsait.png", width=Inches(2.20)
                )
                run = header_para.add_run(" " * 30)
                header_para.add_run().add_picture(
                    "imagens/petrobras.png", width=Inches(2.1)
                )

                # Título centralizado
                title_para = header.add_paragraph()
                title_run = title_para.add_run("\nDocumento de Evidência")
                title_run.bold = True
                title_run.font.size = Pt(20)  # ajusta o tamanho da fonte
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Linha horizontal
                line_para = header.add_paragraph()
                run = line_para.add_run("_" * 100)  # linha horizontal simulada
                run.bold = True
                line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Cria a tabela e define o estilo
                table = doc.add_table(rows=0, cols=2)
                table.style = "Table Grid"

                # adiciona dados da planilha
                for c in row.index:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(c)
                    row_cells[1].text = str(row[c])

                # adiciona dados extras
                for chave, valor in dados.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(chave)
                    row_cells[1].text = str(valor)

                return doc, id_str

            def salvar_documento(base_path, doc, id_str, sistema):
                import os

                # Cria o diretório base se não existir
                os.makedirs(base_path, exist_ok=True)

                nome_arquivo = f"Documento de teste - {sistema} - TC{id_str}.docx"
                full_path = os.path.join(base_path, nome_arquivo)

                doc.save(full_path)
                print(f"Documento gerado: {nome_arquivo}")

            # agora o loop chama as funções internas para cada linha
            for _, row in df_tabela.iterrows():
                doc, id_str = criar_um_documento(df, row)
                salvar_documento(base_path, doc, id_str, row["Sistema"])

        except Exception as e:
            print("Erro ao criar documentos:", e)


    def adicionar_steps(self, base_path):
        try:

            def preparar_dados_steps():
                df = pd.read_excel(self.file_path, sheet_name="Planejamento")
                df.columns = df.columns.str.strip()

                colunas = [
                    "ID",
                    "Sistema",          # se existir
                    "Step Name",
                    "Action",
                    "Expected Result",
                ]

                df_steps = df[colunas].copy()
                df_steps = df_steps.dropna(how="all")
                df_steps = df_steps.apply(
                    lambda col: col.apply(
                        lambda x: x.strip() if isinstance(x, str) else x
                    )
                )
                df_steps["ID"] = df_steps["ID"].ffill()
                df_steps["ID"] = pd.to_numeric(
                    df_steps["ID"], errors="coerce"
                ).astype("Int64")

                return df_steps

            # ————— chama a função e obtém os dados
            df_steps = preparar_dados_steps()

            def editar_docs(base_path, df_steps):
                for doc in os.listdir(base_path):
                    if doc.endswith(".docx"):
                        doc_path = os.path.join(base_path, doc)
                        document = Document(doc_path)

                        # Extrai o ID do nome do arquivo
                        id_str = doc.split("TC")[-1].split(".docx")[0]

                        # Filtra os steps correspondentes ao ID
                        steps = df_steps[df_steps["ID"] == int(id_str)]

                        if not steps.empty:
                            # Insere título "Steps" já com nova página
                            head = document.add_heading("Steps", level=1)
                            head.paragraph_format.page_break_before = True

                            for _, step in steps.iterrows():
                                p = document.add_paragraph()

                                # Step Name em negrito
                                p.add_run(f"{step['Step Name']}: ").bold = True

                                # Action em negrito
                                p.add_run("\nAção: ").bold = True
                                p.add_run(f"{step['Action']}\n")

                                # Expected Result em negrito
                                p.add_run("\n\nResultado Esperado: ").bold = True
                                p.add_run(f"{step['Expected Result']}\n\n")

                            document.save(doc_path)
                            print(f"Steps adicionados ao documento: {doc}")

            # ————— chama editar_docs para realmente processar os arquivos
            editar_docs(base_path, df_steps)

        except Exception as e:
            print("Erro ao adicionar steps:", e)